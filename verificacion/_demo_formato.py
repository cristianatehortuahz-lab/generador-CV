# -*- coding: utf-8 -*-
"""
_demo_formato.py — Ejemplo EJECUTABLE de cómo agregar un formato nuevo.

Sigue el "Patrón A" de docs/hojas-de-vida/08-agregar-nuevo-formato.md al pie de la
letra: un generador que recibe el dict de CVData y devuelve bytes de PDF, reutilizando
los helpers compartidos de cv_format_utils (sin tocar la extracción).

Este archivo demuestra que la guía funciona de punta a punta. NO está registrado en la
API (es solo el paso 1 de la guía); registrarlo es añadir ~3 anclas descritas en el doc.

Uso:
  python scripts_dev/_demo_formato.py "Claudia Talero"
"""
import io, os, sys, tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt
import cv_format_utils as U


def _build_docx(cv_data: dict) -> Document:
    """CV demo de una página: nombre + contacto + primeras publicaciones."""
    doc = Document()

    # Nombre
    p = doc.add_paragraph()
    r = p.add_run(U.clean(cv_data.get("name", "")))
    r.bold = True
    r.font.size = Pt(16)

    # Contacto (reutiliza el helper compartido: email, teléfono, ORCID)
    parts = U.contact_parts(cv_data)
    if parts:
        doc.add_paragraph(" · ".join(parts))

    # Publicaciones agrupadas (reutiliza group_publications + publication_citation)
    pubs = cv_data.get("publications", [])
    if pubs:
        h = doc.add_paragraph()
        hr = h.add_run("Selected Publications")
        hr.bold = True
        hr.font.size = Pt(12)
        grouped = U.group_publications(pubs)
        shown = 0
        for label, items in grouped.items():
            for pub in items:
                cite = U.publication_citation(pub)
                if cite:
                    doc.add_paragraph(cite, style=None)
                    shown += 1
                if shown >= 5:
                    break
            if shown >= 5:
                break
    return doc


def generate_demo_pdf(cv_data: dict) -> bytes:
    """Contrato del Patrón A: dict de CVData -> bytes de PDF."""
    doc = _build_docx(cv_data)
    with tempfile.TemporaryDirectory() as d:
        docx_path = os.path.join(d, "cv.docx")
        pdf_path = os.path.join(d, "cv.pdf")
        doc.save(docx_path)
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        with open(pdf_path, "rb") as f:
            return f.read()


if __name__ == "__main__":
    from cv_generator import CVExtractor
    from dataclasses import asdict
    name = sys.argv[1] if len(sys.argv) > 1 else "Claudia Talero"
    cv = CVExtractor().extract_by_name(name)
    if not cv:
        print("No encontrado:", name); sys.exit(1)
    data = asdict(cv)
    pdf = generate_demo_pdf(data)
    out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                       "Nuevo_entorno_local", "formatos", "DEMO_formato.pdf")
    with open(out, "wb") as f:
        f.write(pdf)
    print("[OK] Demo PDF generado (%d bytes) -> %s" % (len(pdf), out))
