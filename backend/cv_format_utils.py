# -*- coding: utf-8 -*-
"""
cv_format_utils.py -- Utilidades compartidas de mapeo de datos para los
generadores de hoja de vida (Harvard y Europass).

Centraliza el parseo/normalizacion que antes estaba duplicado y fragil en
harvard_cv.py y pdf_filler.py, para que ambos formatos consuman los mismos
datos de forma consistente.
"""

import re
import html as _html
from datetime import datetime

# Nota: python-docx se importa lazy dentro de apply_document_metadata /
# apply_traceability_footer para no imponer la dependencia a los otros
# usuarios de este módulo (extractor, tests, etc.).

_CURRENT_YEAR = datetime.now().year

_INST_KEYWORDS = re.compile(
    r'\b(Universid\w*|University|Institut\w*|Pontificia|Colegio|Escuela|'
    r'Polit\xe9cnic\w*|Hochschule|Universit\xe9\w*|Fachhochschule|College|Academy|'
    r'Akademie|Conservatori\w*|Fundaci\xf3n|School of)\b',
    re.IGNORECASE,
)


def clean(text) -> str:
    """Normaliza texto: unescape de entidades HTML y espacios colapsados."""
    if not text:
        return ""
    return re.sub(r'\s+', ' ', _html.unescape(str(text))).strip()


def parse_education(entry) -> dict:
    """Normaliza una entrada de educacion a {year, degree, institution}.

    Acepta dicts ya estructurados o strings libres del extractor.  El parser
    de strings detecta el anio en cualquier posicion (1940-anio actual+1) y la
    institucion por palabras clave ("Universidad", "University", ...),  dejando
    el resto (con sus comas internas) como grado completo.
    """
    if isinstance(entry, dict):
        return {
            "year": clean(entry.get("year", entry.get("date", ""))),
            "degree": clean(entry.get("degree", entry.get("title", ""))),
            "institution": clean(entry.get("institution", entry.get("name", ""))),
        }

    parts = [clean(p) for p in str(entry).split(",") if clean(p)]
    if not parts:
        return {"year": "", "degree": "", "institution": ""}

    year = ""
    year_idx = None
    for i, p in enumerate(parts):
        m = re.search(r'\b(\d{4})\b', p)
        if m and 1940 <= int(m.group(1)) <= _CURRENT_YEAR + 1:
            year = m.group(1)
            year_idx = i
            break

    if year_idx is not None:
        remaining = [p for j, p in enumerate(parts) if j != year_idx]
        leftover = parts[year_idx].replace(year, "", 1).strip(" ,")
        if leftover:
            remaining.insert(0 if year_idx == 0 else year_idx, leftover)
    else:
        remaining = parts

    inst = ""
    inst_idx = None
    for i in reversed(range(len(remaining))):
        if _INST_KEYWORDS.search(remaining[i]):
            inst_idx = i
            break
    if inst_idx is not None:
        inst = ", ".join(remaining[inst_idx:])
        remaining = remaining[:inst_idx]

    degree = ", ".join(remaining) if remaining else ""
    return {"year": year, "degree": degree, "institution": inst}


def split_institution(institution: str):
    """Separa 'Institucion - Ciudad' en (institucion, ubicacion)."""
    if " - " in institution:
        name, loc = institution.rsplit(" - ", 1)
        return clean(name), clean(loc)
    return clean(institution), ""


def contact_parts(cv: dict) -> list:
    """Partes ordenadas para la linea de contacto (address, email, phone, ORCID)."""
    parts = []
    if cv.get("address"):
        parts.append(clean(cv["address"]))
    if cv.get("email"):
        parts.append(clean(cv["email"]))
    if cv.get("phone"):
        parts.append(clean(cv["phone"]))
    if cv.get("orcid"):
        oid = cv["orcid"].rstrip("/").split("/")[-1]
        parts.append("ORCID: " + oid)
    return parts


_PUB_TYPE_LABELS = {
    "AcademicArticle": "Journal Articles",
    "EditorialArticle": "Editorials",
    "Chapter": "Book Chapters",
    "Book": "Books",
    "ConferencePaper": "Conference Papers",
    "ConferencePoster": "Conference Posters",
    "Letter": "Letters",
    "Comment": "Comments",
}


def _pub_sort_key(pub):
    y = pub.get("year", "") if isinstance(pub, dict) else ""
    try:
        return -int(y)
    except (ValueError, TypeError):
        return 0


def group_publications(pubs: list) -> dict:
    """Agrupa publicaciones por tipo, con re-sort descendente por anio dentro de
    cada grupo."""
    order = ["Journal Articles", "Book Chapters", "Books", "Editorials",
             "Conference Papers", "Conference Posters", "Letters", "Comments"]
    grouped: dict[str, list] = {}
    for pub in pubs:
        if isinstance(pub, dict):
            raw = pub.get("type", "") or "Other"
        else:
            raw = "Other"
        label = _PUB_TYPE_LABELS.get(raw, raw if raw else "Other Publications")
        grouped.setdefault(label, []).append(pub)

    for items in grouped.values():
        items.sort(key=_pub_sort_key)

    ranked = {k: grouped[k] for k in order if k in grouped}
    for k in grouped:
        if k not in ranked:
            ranked[k] = grouped[k]
    return ranked


def publication_citation(pub: dict) -> str:
    """Academic citation: Authors (Year). Title. Journal, Vol(Issue), pp. X-Y. DOI."""
    if not isinstance(pub, dict):
        return clean(pub)
    authors = pub.get("authors", [])
    title = clean(pub.get("title", "")).rstrip(".")
    journal = clean(pub.get("journal", ""))
    volume = clean(pub.get("volume", ""))
    issue = clean(pub.get("issue", ""))
    pages = clean(pub.get("pages", ""))
    year = clean(pub.get("year", ""))
    doi = clean(pub.get("doi", ""))

    parts: list[str] = []
    if authors:
        parts.append("; ".join(authors))
    if year:
        parts.append("(" + year + ")")
    cite_head = " ".join(parts)
    if cite_head:
        cite_head += ". "
    cite = cite_head
    if title:
        cite += '“' + title + '”'
    if journal:
        cite += ". " + journal
    if volume:
        cite += ", " + volume
        if issue:
            cite += "(" + issue + ")"
    # Nota: el número de edición/issue solo se muestra junto al volumen. Sin
    # revista ni volumen (caso de este Solr), un "(40)" suelto se leería como
    # un error, así que se omite aunque esté disponible.
    if pages:
        cite += ", pp. " + pages
    if doi:
        cite += ". DOI: " + doi
    return cite.strip()


def skills_list(cv: dict) -> list:
    """Areas de expertise limpias para la seccion de habilidades."""
    return [clean(a) for a in cv.get("expertise_areas", []) if clean(a)]


def apply_document_metadata(doc, holder_name: str, kind: str = "CV",
                            lang: str = "es-CO"):
    """R2 + R14 — Fija metadatos correctos y el idioma del DOCX.

    Sobrescribe core_properties (Author, Title, Subject, Keywords,
    Comments) para que el PDF resultante identifique al titular real y no
    al autor de la plantilla; y setea el idioma en styles + en cada run
    para que LibreOffice no lo vuelva a en-US al convertir.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    display_name = holder_name or "Investigador"
    if "," in display_name:
        last, first = [x.strip() for x in display_name.split(",", 1)]
        display_name = f"{first} {last}".strip()

    core = doc.core_properties
    core.title = f"Hoja de Vida — {display_name}"
    core.author = display_name
    core.last_modified_by = "HUB-UR"
    core.subject = "Curriculum Vitae académico"
    core.keywords = f"curriculum vitae, {kind}, Universidad del Rosario, {display_name}"
    core.comments = (
        "Documento generado desde HUB-UR (research-hub.urosario.edu.co). "
        "Contenido derivado del perfil VIVO del investigador."
    )
    core.language = lang

    # Aplicar w:lang al estilo Normal y a cada run del documento para que
    # LibreOffice/Word respeten el idioma en el PDF resultante.
    _lang, _territory = (lang.split("-") + [""])[:2]
    lang_tag = lang if _territory else _lang

    style = doc.styles["Normal"].element
    rPr = style.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style.insert(0, rPr)
    existing = rPr.find(qn("w:lang"))
    if existing is not None:
        rPr.remove(existing)
    lang_el = OxmlElement("w:lang")
    lang_el.set(qn("w:val"), lang_tag)
    lang_el.set(qn("w:eastAsia"), lang_tag)
    lang_el.set(qn("w:bidi"), lang_tag)
    rPr.append(lang_el)

    for para in doc.paragraphs:
        for run in para.runs:
            r_rPr = run._element.get_or_add_rPr()
            for old in r_rPr.findall(qn("w:lang")):
                r_rPr.remove(old)
            r_lang = OxmlElement("w:lang")
            r_lang.set(qn("w:val"), lang_tag)
            r_lang.append  # noqa (mantener llamada para evitar lint)
            r_rPr.append(r_lang)


def apply_traceability_footer(doc, font: str = "Calibri"):
    """R12 — Pie de página con trazabilidad del origen del documento."""
    from docx.shared import Pt as _Pt, RGBColor as _RGB
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    stamp = datetime.now().strftime("%Y-%m-%d")
    text = (
        "Documento generado desde HUB-UR (research-hub.urosario.edu.co) "
        f"el {stamp}. Contenido derivado del perfil VIVO del investigador."
    )
    for section in doc.sections:
        footer = section.footer
        # Usar el primer párrafo existente del footer para evitar duplicados
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        # Limpiar runs previos por si el template traía algo
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        para.alignment = _ALIGN.CENTER
        run = para.add_run(text)
        run.font.name = font
        run.font.size = _Pt(7.5)
        run.font.color.rgb = _RGB(0x80, 0x80, 0x80)
        # Reset del rFonts para consistencia entre Word y LibreOffice
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
            rFonts.set(qn(attr), font)


def grant_line(g: dict) -> str:
    """Linea de un proyecto/grant: Title [2023 - 2025]."""
    if not isinstance(g, dict):
        return clean(g)
    title = clean(g.get("title", ""))
    year = clean(g.get("year", ""))
    year_end = clean(g.get("year_end", ""))
    line = '“' + title + '”' if title else ""
    if year:
        date_str = year + " – " + year_end if year_end and year_end != year else year
        line += " [" + date_str + "]"
    return line
