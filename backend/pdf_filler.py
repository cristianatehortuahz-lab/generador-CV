"""
pdf_filler.py — Generador de hoja de vida en el formato oficial Europass.

Replica el layout oficial de europass.europa.eu, verificado contra un PDF
oficial de referencia (no incluido en el paquete por contener datos personales):
  - Banda gris (#F3F3F3) de encabezado con el logo Europass arriba a la
    derecha, el nombre grande en gris medio con regla gris debajo, y la línea
    de campos (etiquetas en negrita) dentro de la banda.
  - Secciones con viñeta-punto gris + título en MAYÚSCULAS casi negro + regla
    gris gruesa.
  - Entradas: línea meta (fechas - LUGAR en mayúsculas) en gris claro,
    organización en negrita subrayada + " - " + grado/rol regular en la misma
    línea (sin cursiva), descripción en gris.
  - Paleta 100% gris (el único color es el logo).

Para investigadores se añaden, en el mismo estilo, secciones de producción
académica (Publications, Research Projects) además de Education & Training,
Work Experience y Skills.

Pipeline: python-docx -> DOCX -> docx2pdf (Word) -> PDF bytes.
"""

import os

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

import cv_format_utils as U

def _detect_font():
    """Use Open Sans if installed (official Europass font), fallback to Arial."""
    try:
        import ctypes
        from ctypes import wintypes
        gdi32 = ctypes.windll.gdi32
        hdc = ctypes.windll.user32.GetDC(0)
        LOGFONT = type('LOGFONT', (ctypes.Structure,), {'_fields_': [
            ('lfHeight', ctypes.c_long), ('lfWidth', ctypes.c_long),
            ('lfEscapement', ctypes.c_long), ('lfOrientation', ctypes.c_long),
            ('lfWeight', ctypes.c_long), ('lfItalic', ctypes.c_byte),
            ('lfUnderline', ctypes.c_byte), ('lfStrikeOut', ctypes.c_byte),
            ('lfCharSet', ctypes.c_byte), ('lfOutPrecision', ctypes.c_byte),
            ('lfClipPrecision', ctypes.c_byte), ('lfQuality', ctypes.c_byte),
            ('lfPitchAndFamily', ctypes.c_byte),
            ('lfFaceName', ctypes.c_char * 32),
        ]})
        found = [False]
        @ctypes.WINFUNCTYPE(ctypes.c_int, ctypes.POINTER(LOGFONT),
                            ctypes.c_void_p, wintypes.DWORD, ctypes.c_void_p)
        def cb(lf, tm, ft, lp):
            found[0] = True
            return 0
        lf = LOGFONT()
        lf.lfFaceName = b"Open Sans"
        lf.lfCharSet = 0
        gdi32.EnumFontFamiliesExA(hdc, ctypes.byref(lf), cb, 0, 0)
        ctypes.windll.user32.ReleaseDC(0, hdc)
        return "Open Sans" if found[0] else "Arial"
    except Exception:
        return "Arial"

FONT = _detect_font()

# Paleta del Europass oficial (medida sobre el PDF de referencia; todo gris,
# el único color va en el logo).
# Colores medidos del Europass oficial (PDF de referencia europass.europa.eu):
#   nombre/org/meta/reglas = #4F4F4F · títulos y etiquetas = #000000 (negro)
#   cuerpo = #5A5959 · banda = #F4F4F4
BAND_FILL = "F4F4F4"                        # banda superior full-bleed (oficial)
NAME_GREY = RGBColor(0x4F, 0x4F, 0x4F)      # nombre (oficial: regular #4F4F4F)
TITLE_DARK = RGBColor(0x00, 0x00, 0x00)     # títulos de sección y etiquetas (negro)
TEXT_GREY = RGBColor(0x5A, 0x59, 0x59)      # cuerpo/descripción (oficial #5A5959)
META_GREY = RGBColor(0x4F, 0x4F, 0x4F)      # línea meta/fechas (oficial #4F4F4F)
DOT_GREY = RGBColor(0x4F, 0x4F, 0x4F)       # cuadro de viñeta de sección
RULE_GREY = "4F4F4F"                        # reglas horizontales (oficial #4F4F4F)

# Geometría A4 (pt) y banda superior, medidas del Europass oficial.
PAGE_W_PT = 595
BAND_H_PT = 134

_LOGO = os.environ.get("EUROPASS_LOGO", "").strip() or os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "formatos", "europass_logo.png",
)

MAX_PUBS_PER_TYPE = 200  # tope alto: lista prácticamente todo (evita que el conteo del encabezado mienta)


def _sf(run, size=Pt(10), bold=False, underline=False, color=TEXT_GREY, font=FONT):
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    run.font.underline = underline
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rFonts.set(qn(attr), font)


def _rule(paragraph, color=RULE_GREY, sz="12"):
    """Regla horizontal inferior en un párrafo."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _eu_item(doc, text, size=Pt(9), color=TEXT_GREY):
    """Item de lista Europass con viñeta y sangría francesa.

    La viñeta queda a la izquierda y las líneas que envuelven se alinean bajo
    el texto (no bajo la viñeta), separando visualmente cada publicación/proyecto.
    """
    pb = doc.add_paragraph()
    pf = pb.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.left_indent = Cm(0.7)
    pf.first_line_indent = Cm(-0.3)  # sangría francesa: la viñeta sobresale
    _sf(pb.add_run("•  "), size=size, color=color)
    _sf(pb.add_run(text), size=size, color=color)
    return pb


def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        borders.append(el)
    tblPr.append(borders)


def _band_shape(paragraph, w_pt=PAGE_W_PT, h_pt=BAND_H_PT, fill=BAND_FILL):
    """Inserta un rectángulo flotante full-bleed anclado a la esquina superior
    izquierda de la PÁGINA (detrás del texto), reproduciendo la banda gris
    del Europass oficial que cubre todo el ancho superior."""
    style = (
        "position:absolute;margin-left:0;margin-top:0;"
        f"width:{w_pt}pt;height:{h_pt}pt;z-index:-251658240;"
        "mso-position-horizontal-relative:page;"
        "mso-position-vertical-relative:page"
    )
    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word">'
        '<w:pict>'
        f'<v:rect style="{style}" fillcolor="#{fill}" stroked="f">'
        '<w10:wrap type="none"/></v:rect>'
        '</w:pict></w:r>'
    )
    paragraph._p.append(parse_xml(xml))


def _header(doc, name, cv):
    """Encabezado Europass oficial: banda gris full-bleed (toca los bordes
    superior/izquierdo/derecho), logo arriba-derecha, nombre con regla y línea
    de campos, todo sobre la banda."""
    # Párrafo que porta la banda flotante + el logo alineado a la derecha.
    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_top.paragraph_format.space_before = Pt(0)
    p_top.paragraph_format.space_after = Pt(2)
    _band_shape(p_top)
    if os.path.exists(_LOGO):
        try:
            p_top.add_run().add_picture(_LOGO, width=Inches(1.5))
        except Exception:
            pass

    # Nombre con regla gris debajo.
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(2)
    p_name.paragraph_format.space_after = Pt(6)
    _sf(p_name.add_run(name), size=Pt(16), bold=False, color=NAME_GREY)
    _rule(p_name, color=RULE_GREY, sz="6")

    # Línea de campos con etiquetas en negrita (omitiendo los ausentes).
    fields = []
    if cv.get("nationality"):
        fields.append(("Nationality", U.clean(cv["nationality"])))
    if cv.get("address"):
        fields.append(("Address", U.clean(cv["address"])))
    if cv.get("email"):
        fields.append(("Email", U.clean(cv["email"])))
    if cv.get("phone"):
        fields.append(("Phone", U.clean(cv["phone"])))
    if cv.get("orcid"):
        oid = cv["orcid"].rstrip("/").split("/")[-1]
        fields.append(("ORCID", oid))
    if fields:
        p_f = doc.add_paragraph()
        p_f.paragraph_format.space_before = Pt(2)
        p_f.paragraph_format.space_after = Pt(0)
        for i, (label, value) in enumerate(fields):
            if i > 0:
                _sf(p_f.add_run("  |  "), size=Pt(10), color=META_GREY)
            _sf(p_f.add_run(f"{label}: "), size=Pt(10), bold=False, color=TITLE_DARK)
            _sf(p_f.add_run(value), size=Pt(10), color=TEXT_GREY)

    # Espaciador para bajar el contenido por debajo de la banda (134pt).
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(18)


def _eu_section(doc, title):
    """Sección oficial: viñeta-punto gris + MAYÚSCULAS casi negro + regla gruesa."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    # Cuadro de viñeta colgado hacia el margen izquierdo (como el oficial).
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    _sf(p.add_run("▪"), size=Pt(11), color=DOT_GREY)
    _sf(p.add_run("  "), size=Pt(11))
    _sf(p.add_run(title.upper()), size=Pt(11), bold=False, color=TITLE_DARK)
    _rule(p, color=RULE_GREY, sz="11")
    return p


def _entry(doc, meta="", org="", subtitle="", desc="", extra_label="", extra_value=""):
    """Entrada oficial Europass:
      meta (fechas - LUGAR) gris claro
      **Org** (negrita, subrayada) - subtítulo regular, en la MISMA línea
      descripción gris
      Etiqueta extra: valor (p. ej. Level in EQF)."""
    if meta:
        pm = doc.add_paragraph()
        pm.paragraph_format.space_before = Pt(6)
        pm.paragraph_format.space_after = Pt(1)
        _sf(pm.add_run(meta), size=Pt(9), color=META_GREY)
    if org or subtitle:
        po = doc.add_paragraph()
        po.paragraph_format.space_before = Pt(0)
        po.paragraph_format.space_after = Pt(2)
        if org:
            _sf(po.add_run(org), size=Pt(11), bold=False,
                underline=False, color=TEXT_GREY)
        if subtitle:
            if org:
                _sf(po.add_run(" - "), size=Pt(11), color=TEXT_GREY)
            _sf(po.add_run(subtitle), size=Pt(11), color=TEXT_GREY)
    if desc:
        pd = doc.add_paragraph()
        pd.paragraph_format.space_before = Pt(2)
        pd.paragraph_format.space_after = Pt(2)
        _sf(pd.add_run(desc), size=Pt(10), color=TEXT_GREY)
    if extra_label:
        pe = doc.add_paragraph()
        pe.paragraph_format.space_before = Pt(0)
        pe.paragraph_format.space_after = Pt(2)
        _sf(pe.add_run(f"{extra_label}: "), size=Pt(9), bold=False, color=TITLE_DARK)
        _sf(pe.add_run(extra_value), size=Pt(9), color=TEXT_GREY)


def _build_docx(cv_data: dict) -> Document:
    """Construye la hoja de vida Europass con datos reales del investigador."""
    doc = Document()

    # A4 con márgenes tipo Europass.
    s = doc.sections[0]
    s.page_width = Cm(21.0)   # A4 width
    s.page_height = Cm(29.7)  # A4 height
    s.top_margin = Cm(0.85)
    s.bottom_margin = Cm(1.3)
    s.left_margin = Cm(1.38)
    s.right_margin = Cm(1.82)

    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(10)

    # Nombre: "Apellidos, Nombres" -> "Nombres Apellidos".
    name = U.clean(cv_data.get("name", ""))
    if not name:
        # R10 — sin nombre no se puede generar un CV presentable.
        raise ValueError("cv_data.name está vacío; el CV Europass requiere un titular.")
    if "," in name:
        last, first = [x.strip() for x in name.split(",", 1)]
        name = f"{first} {last}".strip()

    _header(doc, name, cv_data)

    # About Me (overview) — campo soportado por Europass.
    overview = U.clean(cv_data.get("overview", ""))
    if overview:
        _eu_section(doc, "About Me")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        _sf(p.add_run(overview), size=Pt(10), color=TEXT_GREY)

    # Education & Training
    edu_list = cv_data.get("education", [])
    if edu_list:
        _eu_section(doc, "Education & Training")
        for edu in edu_list:
            e = U.parse_education(edu)
            inst, loc = U.split_institution(e["institution"])
            meta = e["year"]
            if loc:
                # El oficial muestra el lugar en MAYÚSCULAS tras las fechas.
                meta = f"{meta}  -  {loc.upper()}" if meta else loc.upper()
            _entry(doc, meta=meta, org=inst or e["degree"],
                   subtitle=e["degree"] if inst else "")

    # Work Experience
    job = U.clean(cv_data.get("job_title", ""))
    org = U.clean(cv_data.get("organization", ""))
    affs = cv_data.get("affiliations", [])
    if job or affs:
        _eu_section(doc, "Work Experience")
        if job:
            _entry(doc, org=org or "Universidad del Rosario", subtitle=job)
        for a in affs:
            if isinstance(a, dict):
                a_name = U.clean(a.get("name", a.get("organization", "")))
                role = U.clean(a.get("role", a.get("position", "")))
                if a_name:
                    _entry(doc, org=a_name, subtitle=role)

    # Publications (agrupadas por tipo)
    pubs = cv_data.get("publications", [])
    if pubs:
        grouped = U.group_publications(pubs)
        _eu_section(doc, "Publications")
        for label, items in grouped.items():
            ps = doc.add_paragraph()
            ps.paragraph_format.space_before = Pt(6)
            ps.paragraph_format.space_after = Pt(2)
            _sf(ps.add_run(f"{label} ({len(items)})"), size=Pt(10.5),
                bold=True, color=TITLE_DARK)
            for pub in items[:MAX_PUBS_PER_TYPE]:
                cite = U.publication_citation(pub)
                if cite:
                    _eu_item(doc, cite, size=Pt(9))

    # Research Projects (grants)
    grants = cv_data.get("grants", [])
    if grants:
        _eu_section(doc, "Research Projects")
        for g in grants:
            line = U.grant_line(g)
            if line:
                _eu_item(doc, line, size=Pt(10))

    # Advised Theses (tesis dirigidas, listadas individualmente)
    theses = cv_data.get("theses", [])
    if theses:
        _eu_section(doc, "Advised Theses")
        for t in theses:
            thesis = t if isinstance(t, dict) else {"title": str(t)}
            title = U.clean(thesis.get("title", ""))
            if not title:
                continue
            year = U.clean(thesis.get("year", ""))
            line = f"{title} [{year}]" if year else title
            _eu_item(doc, line, size=Pt(10))

    # Skills (áreas de expertise separadas por |)
    skills = U.skills_list(cv_data)
    if skills:
        _eu_section(doc, "Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        _sf(p.add_run("  |  ".join(skills)), size=Pt(10), color=TEXT_GREY)

    # Pie de trazabilidad (R12)
    U.apply_traceability_footer(doc, font=FONT)

    # Metadatos, idioma (R2, R14)
    U.apply_document_metadata(doc, name, kind="Europass", lang="es-CO")

    return doc


def fill_europass(cv_data: dict, template_path: str, output_path: str):
    """Genera la hoja de vida Europass como PDF (template_path es solo gate de
    disponibilidad; el documento se construye programáticamente).
    
    Detecta el sistema operativo y usa docx2pdf en Windows o LibreOffice headless en Linux.
    """
    import platform
    import subprocess
    
    docx_path = output_path.replace(".pdf", ".docx")
    doc = _build_docx(cv_data)
    doc.save(docx_path)
    
    try:
        if platform.system() == "Windows":
            import pythoncom
            pythoncom.CoInitialize()
            try:
                try:
                    from docx2pdf import convert
                except ImportError as e:
                    raise RuntimeError(
                        "docx2pdf no está instalado; es necesario para generar el PDF."
                    ) from e
                try:
                    convert(docx_path, output_path)
                except Exception as e:
                    raise RuntimeError(
                        "No se pudo convertir el DOCX a PDF. Verifique que Microsoft Word "
                        "esté instalado y disponible (docx2pdf lo usa vía COM)."
                    ) from e
            finally:
                pythoncom.CoUninitialize()
        else:
            # En Linux usar LibreOffice headless
            soffice = os.environ.get("SOFFICE_PATH", "soffice")
            outdir = os.path.dirname(output_path) or "."
            try:
                subprocess.run([
                    soffice, "--headless", "--convert-to", "pdf",
                    "--outdir", outdir, docx_path
                ], check=True, timeout=120)
            except Exception as e:
                raise RuntimeError(
                    f"Error convirtiendo DOCX a PDF usando LibreOffice ({soffice}): {e}"
                ) from e
            
            # LibreOffice escribe <nombre_archivo>.pdf en outdir
            expected = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if expected != output_path and os.path.exists(expected):
                if os.path.exists(output_path):
                    try:
                        os.remove(output_path)
                    except OSError:
                        pass
                os.rename(expected, output_path)
                
        print(f"[OK] Europass PDF -> {output_path}")
    finally:
        if os.path.exists(docx_path):
            try:
                os.remove(docx_path)
            except OSError:
                pass

