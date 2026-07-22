"""
harvard_cv.py — Generador de hoja de vida en el formato oficial de Harvard.

Clona la plantilla oficial de Harvard College
(`formatos/2025-template_bullet (3).docx`), verificada contra su render real:
carta 8.5x11", **Calibri 11pt** en todo el documento (los runs de la plantilla
traen Calibri directo), nombre centrado en negrita con una única línea
horizontal debajo, contacto centrado con separador "•", títulos de sección
**centrados en negrita sin bordes**, entradas con organización en negrita +
ubicación/fechas alineadas a la derecha (tab derecho) y viñetas de logros.

Para investigadores se añaden, en el MISMO estilo oficial, secciones de
producción académica (Publications, Research Projects) además de las estándar
(Education, Experience, Skills & Interests).

Pipeline: python-docx -> DOCX -> docx2pdf (Word) -> PDF bytes.
"""

import os
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import cv_format_utils as U

FONT = "Calibri"
BLACK = RGBColor(0, 0, 0)

# Plantilla oficial (configurable por entorno).
_TEMPLATE = os.environ.get("HARVARD_TEMPLATE", "").strip() or os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Nuevo_entorno_local", "formatos", "2025-template_bullet (3).docx",
)

# Máximo de publicaciones por tipo, para no desbordar el documento.
MAX_PUBS_PER_TYPE = 200  # tope alto: lista prácticamente todo (evita que el conteo del encabezado mienta)


def _sf(run, size=Pt(11), bold=False, italic=False, color=BLACK, font=FONT):
    """Fija la fuente de un run (incluye eastAsia para consistencia en Word)."""
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rFonts.set(qn(attr), font)


def _name_rule(paragraph):
    """Línea horizontal bajo el nombre — la ÚNICA línea del formato oficial
    (en la plantilla es una forma VML de 1pt a lo ancho de la página)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')  # 1pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _clear_body(doc):
    """Elimina todos los párrafos/tablas del cuerpo, conservando la geometría
    (sectPr con tamaño de página y márgenes de la plantilla)."""
    body = doc.element.body
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)


def _right_tab_pos(doc):
    """Posición del margen derecho (para tab derecho de ubicación/fechas)."""
    s = doc.sections[0]
    return s.page_width - s.left_margin - s.right_margin


def _load_template():
    """Carga la plantilla oficial (o un documento en blanco equivalente)."""
    if os.path.exists(_TEMPLATE):
        doc = Document(_TEMPLATE)
        _clear_body(doc)
    else:
        doc = Document()
        s = doc.sections[0]
        s.page_width = Inches(8.5)
        s.page_height = Inches(11)
        s.left_margin = s.right_margin = Inches(0.42)
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.19)
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(11)
    return doc


def _section(doc, title):
    """Título de sección estilo oficial: CENTRADO, negrita, sin bordes."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    _sf(p.add_run(title), size=Pt(11), bold=True)
    return p


def _entry(doc, left, right="", line2_left="", line2_right="",
           line2_bold=False, bullets=None, right_tab=Inches(7.66)):
    """Entrada oficial: `Organización(negrita)⇥Ubicación` y segunda línea
    `Cargo/Grado⇥Fechas` (negrita solo en experiencia), más viñetas."""
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)
    _sf(p1.add_run(left), size=Pt(11), bold=True)
    if right:
        _sf(p1.add_run("\t"), size=Pt(11))
        _sf(p1.add_run(right), size=Pt(11))

    if line2_left or line2_right:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(1)
        p2.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)
        if line2_left:
            _sf(p2.add_run(line2_left), size=Pt(11), bold=line2_bold)
        if line2_right:
            _sf(p2.add_run("\t"), size=Pt(11))
            _sf(p2.add_run(line2_right), size=Pt(11))

    for b in (bullets or []):
        _bullet(doc, b)


def _bullet(doc, text, size=Pt(11)):
    prefix = ""
    try:
        pb = doc.add_paragraph(style='List Bullet')
    except KeyError:
        # Sin estilo de lista en la plantilla: viñeta manual.
        pb = doc.add_paragraph()
        prefix = "•  "
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after = Pt(1)
    pb.paragraph_format.left_indent = Inches(0.55)
    _sf(pb.add_run(prefix + text), size=size)


def _build_docx(cv_data: dict) -> Document:
    """Construye la hoja de vida Harvard con datos reales del investigador."""
    doc = _load_template()
    rtab = _right_tab_pos(doc)

    # ── Nombre (centrado, negrita, con la línea oficial debajo) ────────
    name = U.clean(cv_data.get("name", ""))
    if not name:
        # R10 — sin nombre no se puede generar un CV presentable. El endpoint
        # ya devuelve 422 antes de llegar aquí; este raise es una red de seguridad.
        raise ValueError("cv_data.name está vacío; el CV Harvard requiere un titular.")
    # El extractor entrega "Apellidos, Nombres"; mostrar "Nombres Apellidos".
    if "," in name:
        last, first = [x.strip() for x in name.split(",", 1)]
        name = f"{first} {last}".strip()
    pn = doc.add_paragraph()
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pn.paragraph_format.space_after = Pt(6)
    _sf(pn.add_run(name), size=Pt(11), bold=True)
    _name_rule(pn)

    # ── Contacto (centrado con "•") ─────────────────────────────────────
    parts = U.contact_parts(cv_data)
    if parts:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before = Pt(4)
        pc.paragraph_format.space_after = Pt(4)
        for i, part in enumerate(parts):
            if i > 0:
                _sf(pc.add_run(" • "), size=Pt(11))
            _sf(pc.add_run(part), size=Pt(11))

    # ── Education ──────────────────────────────────────────────────────
    edu_list = cv_data.get("education", [])
    if edu_list:
        _section(doc, "Education")
        for edu in edu_list:
            e = U.parse_education(edu)
            inst, loc = U.split_institution(e["institution"])
            _entry(doc, inst or e["degree"], loc,
                   line2_left=e["degree"] if inst else "",
                   line2_right=e["year"], right_tab=rtab)

    # ── Experience ─────────────────────────────────────────────────────
    job = U.clean(cv_data.get("job_title", ""))
    org = U.clean(cv_data.get("organization", ""))
    affs = cv_data.get("affiliations", [])
    if job or affs:
        _section(doc, "Experience")
        if job:
            _entry(doc, org or "Universidad del Rosario", "",
                   line2_left=job, line2_bold=True, right_tab=rtab)
        for a in affs:
            if isinstance(a, dict):
                a_name = U.clean(a.get("name", a.get("organization", "")))
                role = U.clean(a.get("role", a.get("position", "")))
                if a_name:
                    _entry(doc, a_name, "", line2_left=role,
                           line2_bold=True, right_tab=rtab)
            elif a:
                _entry(doc, U.clean(a), right_tab=rtab)

    # ── Publications (agrupadas por tipo) ──────────────────────────────
    pubs = cv_data.get("publications", [])
    if pubs:
        grouped = U.group_publications(pubs)
        for label, items in grouped.items():
            _section(doc, f"{label} ({len(items)})")
            for pub in items[:MAX_PUBS_PER_TYPE]:
                cite = U.publication_citation(pub)
                if cite:
                    _bullet(doc, cite, size=Pt(10.5))

    # ── Research Projects (grants) ─────────────────────────────────────
    grants = cv_data.get("grants", [])
    if grants:
        _section(doc, "Research Projects")
        for g in grants:
            line = U.grant_line(g)
            if line:
                _bullet(doc, line)

    # ── Advised Theses ─────────────────────────────────────────────────
    theses = cv_data.get("theses", [])
    if theses:
        _section(doc, "Advised Theses")
        for t in theses:
            thesis = t if isinstance(t, dict) else {"title": str(t)}
            title = thesis.get("title", "")
            year = thesis.get("year", "")
            line = f"{title} [{year}]" if year else title
            _bullet(doc, line)

    # ── Skills & Interests ─────────────────────────────────────────────
    skills = U.skills_list(cv_data)
    if skills:
        _section(doc, "Skills & Interests")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        _sf(p.add_run("Research Areas: "), size=Pt(11), bold=True)
        _sf(p.add_run(", ".join(skills)), size=Pt(11))

    # ── Pie de trazabilidad (R12) ──────────────────────────────────────
    U.apply_traceability_footer(doc, font=FONT)

    # ── Metadatos, idioma (R2, R14) ────────────────────────────────────
    U.apply_document_metadata(doc, name, kind="Harvard", lang="es-CO")

    return doc


def generate_harvard_pdf(cv_data: dict) -> bytes:
    """Genera la hoja de vida Harvard como bytes PDF.
    
    Detecta el sistema operativo y usa docx2pdf en Windows o LibreOffice headless en Linux.
    """
    import platform
    import subprocess
    
    doc = _build_docx(cv_data)
    with tempfile.TemporaryDirectory() as tmpdir:
        dx = os.path.join(tmpdir, "harvard_cv.docx")
        px = os.path.join(tmpdir, "harvard_cv.pdf")
        doc.save(dx)
        
        if platform.system() == "Windows":
            import pythoncom
            pythoncom.CoInitialize()
            try:
                try:
                    from docx2pdf import convert
                except ImportError as e:
                    raise RuntimeError(
                        "docx2pdf no está instalado; es necesario para generar el PDF."
                    ) from e
                try:
                    convert(dx, px)
                except Exception as e:
                    raise RuntimeError(
                        "No se pudo convertir el DOCX a PDF. Verifique que Microsoft Word "
                        "esté instalado y disponible (docx2pdf lo usa vía COM)."
                    ) from e
            finally:
                pythoncom.CoUninitialize()
        else:
            # En Linux usar LibreOffice headless
            soffice = os.environ.get("SOFFICE_PATH", "soffice")
            try:
                subprocess.run([
                    soffice, "--headless", "--convert-to", "pdf",
                    "--outdir", tmpdir, dx
                ], check=True, timeout=120)
            except Exception as e:
                raise RuntimeError(
                    f"Error convirtiendo DOCX a PDF usando LibreOffice ({soffice}): {e}"
                ) from e
            
            # LibreOffice escribe harvard_cv.pdf en el outdir
            expected = os.path.join(tmpdir, "harvard_cv.pdf")
            if not os.path.exists(expected):
                files = [f for f in os.listdir(tmpdir) if f.endswith(".pdf")]
                if files:
                    expected = os.path.join(tmpdir, files[0])
            px = expected
            
        with open(px, "rb") as f:
            return f.read()

